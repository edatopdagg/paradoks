import hashlib
import io
import re
import shutil
import subprocess
import tempfile
import zipfile
from dataclasses import dataclass
from pathlib import Path

import requests
from bs4 import BeautifulSoup
from docx import Document

from v3_identity import (
    VersionIdentity,
    infer_version_identity,
)


HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 "
        "(compatible; ParadoksV3/1.0)"
    )
}

TIMEOUT_SECONDS = 120


@dataclass(frozen=True)
class FetchedDocument:
    org: str
    code: str
    title: str
    version: str
    release: str
    source_url: str
    source_filename: str
    local_path: str
    package_path: str
    extracted_text_path: str
    content_sha256: str
    document_text: str
    page_texts: tuple[str, ...]


def _slug(value: str) -> str:
    return re.sub(
        r"[^a-z0-9]+",
        "-",
        value.casefold(),
    ).strip("-")


def _sha256_bytes(
    value: bytes,
) -> str:
    return hashlib.sha256(
        value
    ).hexdigest()


def _download(
    source_url: str,
) -> tuple[bytes, str]:
    response = requests.get(
        source_url,
        headers=HEADERS,
        timeout=TIMEOUT_SECONDS,
    )

    response.raise_for_status()

    return (
        response.content,
        response.headers.get(
            "Content-Type",
            "",
        ),
    )


def _read_docx_bytes(
    raw_bytes: bytes,
) -> str:
    document = Document(
        io.BytesIO(raw_bytes)
    )

    parts: list[str] = []

    if hasattr(
        document,
        "iter_inner_content",
    ):
        blocks = (
            document.iter_inner_content()
        )
    else:
        blocks = (
            list(document.paragraphs)
            + list(document.tables)
        )

    for block in blocks:
        if hasattr(block, "rows"):
            for row in block.rows:
                cells = [
                    " ".join(
                        cell.text.split()
                    )
                    for cell in row.cells
                    if cell.text.strip()
                ]

                if cells:
                    parts.append(
                        " | ".join(cells)
                    )

        else:
            text = " ".join(
                (
                    getattr(
                        block,
                        "text",
                        "",
                    )
                    or ""
                ).split()
            )

            if text:
                parts.append(text)

    return "\n".join(parts)


def _read_legacy_doc_bytes(
    raw_bytes: bytes,
) -> str:
    antiword_path = shutil.which(
        "antiword"
    )

    if antiword_path is None:
        raise RuntimeError(
            "Legacy .doc belgesi bulundu ancak "
            "antiword kurulu de?il."
        )

    with tempfile.TemporaryDirectory(
        prefix="paradoks_legacy_doc_"
    ) as temporary_directory:
        document_path = (
            Path(temporary_directory)
            / "document.doc"
        )

        document_path.write_bytes(
            raw_bytes
        )

        result = subprocess.run(
            [
                antiword_path,
                "-m",
                "UTF-8.txt",
                str(document_path),
            ],
            capture_output=True,
            check=False,
            timeout=120,
        )

    if result.returncode != 0:
        error_message = result.stderr.decode(
            "utf-8",
            errors="replace",
        ).strip()

        raise ValueError(
            "Legacy .doc metni okunamad?: "
            f"{error_message}"
        )

    return result.stdout.decode(
        "utf-8",
        errors="replace",
    ).strip()


def _extract_3gpp_docx_parts(
    raw_zip: bytes,
    requested_code: str,
) -> list[tuple[str, bytes]]:
    requested_digits = "".join(
        re.findall(
            r"\d",
            requested_code,
        )
    )

    with zipfile.ZipFile(
        io.BytesIO(raw_zip)
    ) as archive:
        candidates = [
            name
            for name in archive.namelist()
            if (
                name.casefold().endswith(
                    (".docx", ".doc")
                )
                and not name.startswith(
                    "__MACOSX/"
                )
                and not Path(
                    name
                ).name.startswith(
                    ("~$", "._")
                )
            )
        ]

        matching = [
            name
            for name in candidates
            if requested_digits
            in Path(name).name
        ]

        selected = (
            matching or candidates
        )

        if not selected:
            raise ValueError(
                "3GPP ZIP paketi içinde "
                "DOCX bulunamadı."
            )

        def part_order(
            value: str,
        ) -> tuple[int, str]:
            filename = Path(
                value
            ).name

            match = re.search(
                r"_(\d+)(?:_|$)",
                Path(filename).stem,
            )

            number = (
                int(match.group(1))
                if match
                else 0
            )

            return (
                number,
                filename.casefold(),
            )

        selected.sort(
            key=part_order
        )

        return [
            (
                Path(name).name,
                archive.read(name),
            )
            for name in selected
        ]


def _read_html(
    raw_bytes: bytes,
) -> str:
    soup = BeautifulSoup(
        raw_bytes,
        "html.parser",
    )

    for element in soup(
        [
            "script",
            "style",
            "nav",
            "footer",
        ]
    ):
        element.decompose()

    root = (
        soup.find("main")
        or soup.find("article")
        or soup.body
        or soup
    )

    lines = [
        " ".join(line.split())
        for line
        in root.get_text(
            "\n"
        ).splitlines()
    ]

    return "\n".join(
        line
        for line in lines
        if line
    )


def _read_pdf(
    raw_bytes: bytes,
) -> tuple[str, tuple[str, ...]]:
    import pdfplumber

    pages: list[str] = []

    with pdfplumber.open(
        io.BytesIO(raw_bytes)
    ) as pdf:
        for page in pdf.pages:
            pages.append(
                (
                    page.extract_text()
                    or ""
                ).strip()
            )

    marked_text = "\n\n".join(
        (
            f"[[PAGE:{index}]]\n"
            f"{text}"
        )
        for index, text in enumerate(
            pages,
            start=1,
        )
    )

    return (
        marked_text,
        tuple(pages),
    )


def _validate_identity(
    *,
    org: str,
    code: str,
    identity: VersionIdentity,
    document_text: str,
) -> None:
    head = "\n".join(
        document_text
        .splitlines()[:120]
    )

    organization = org.upper()

    if organization == "3GPP":
        match = re.search(
            (
                r"3GPP\s+(TS|TR)\s+"
                r"(\d{2}\.\d{3})\s+"
                r"V(\d+\.\d+\.\d+)"
            ),
            head,
            flags=re.IGNORECASE,
        )

        if not match:
            raise ValueError(
                "3GPP belge başlığından "
                "kimlik okunamadı."
            )

        detected_code = (
            f"{match.group(1).upper()} "
            f"{match.group(2)}"
        )

        if (
            detected_code.casefold()
            != code.casefold()
        ):
            raise ValueError(
                "Yanlış 3GPP belgesi: "
                f"{detected_code} != {code}"
            )

        if (
            match.group(3)
            != identity.version
        ):
            raise ValueError(
                "3GPP sürüm uyuşmazlığı: "
                f"{match.group(3)} != "
                f"{identity.version}"
            )

    elif organization == "ETSI":
        number = "".join(
            re.findall(
                r"\d",
                code,
            )
        )

        head_digits = "".join(
            re.findall(
                r"\d",
                head,
            )
        )

        if number not in head_digits:
            raise ValueError(
                "ETSI belge kodu içerikte "
                f"doğrulanamadı: {code}"
            )

        version_match = re.search(
            r"V(\d+\.\d+\.\d+)",
            head,
            re.IGNORECASE,
        )

        if (
            version_match
            and version_match.group(1)
            != identity.version
        ):
            raise ValueError(
                "ETSI sürüm uyuşmazlığı: "
                f"{version_match.group(1)} "
                f"!= {identity.version}"
            )

    elif organization == "IETF":
        number_match = re.search(
            r"(\d{3,5})",
            code,
        )

        if (
            not number_match
            or not re.search(
                (
                    rf"\bRFC\s*"
                    rf"{re.escape(number_match.group(1))}"
                    rf"\b"
                ),
                head,
                flags=re.IGNORECASE,
            )
        ):
            raise ValueError(
                "RFC kimliği içerikte "
                f"doğrulanamadı: {code}"
            )


def fetch_document(
    *,
    org: str,
    code: str,
    title: str,
    source_url: str,
    output_root: str | Path,
) -> FetchedDocument:
    identity = infer_version_identity(
        org=org,
        code=code,
        source_url=source_url,
    )

    (
        raw_bytes,
        content_type,
    ) = _download(source_url)

    output_root = Path(
        output_root
    ).resolve()

    directory = (
        output_root
        / "documents"
        / _slug(org)
        / _slug(code)
        / _slug(
            identity.version
            or "unknown"
        )
    )

    directory.mkdir(
        parents=True,
        exist_ok=True,
    )

    package_path = (
        directory
        / identity.source_filename
    )

    package_path.write_bytes(
        raw_bytes
    )

    primary_path = package_path

    page_texts: tuple[
        str,
        ...,
    ] = ()

    lower_url = (
        source_url.casefold()
    )

    if (
        org.upper() == "3GPP"
        or lower_url.endswith(".zip")
    ):
        parts = (
            _extract_3gpp_docx_parts(
                raw_bytes,
                code,
            )
        )

        stored_parts: list[Path] = []
        part_texts: list[str] = []

        for (
            docx_name,
            docx_bytes,
        ) in parts:
            part_path = (
                directory / docx_name
            )

            part_path.write_bytes(
                docx_bytes
            )

            stored_parts.append(
                part_path
            )

            if docx_name.casefold().endswith(
                ".docx"
            ):
                part_text = _read_docx_bytes(
                    docx_bytes
                )
            else:
                part_text = (
                    _read_legacy_doc_bytes(
                        docx_bytes
                    )
                )

            part_texts.append(
                part_text
            )

        document_text = "\n".join(
            text
            for text in part_texts
            if text
        )

        if len(stored_parts) == 1:
            primary_path = (
                stored_parts[0]
            )
        else:
            # Çok parçalı belgede bütün kaynağı
            # temsil eden resmi ZIP kullanılır.
            primary_path = package_path

    elif (
        lower_url.endswith(".pdf")
        or "application/pdf"
        in content_type.casefold()
    ):
        (
            document_text,
            page_texts,
        ) = _read_pdf(raw_bytes)

    else:
        document_text = (
            _read_html(raw_bytes)
        )

    _validate_identity(
        org=org,
        code=code,
        identity=identity,
        document_text=document_text,
    )

    extracted_text_path = (
        directory / "extracted.txt"
    )

    extracted_text_path.write_text(
        document_text,
        encoding="utf-8",
    )

    return FetchedDocument(
        org=org.upper(),
        code=code.upper(),
        title=title.strip(),
        version=identity.version,
        release=identity.release,
        source_url=source_url,
        source_filename=(
            identity.source_filename
        ),
        local_path=(
            primary_path
            .relative_to(output_root)
            .as_posix()
        ),
        package_path=(
            package_path
            .relative_to(output_root)
            .as_posix()
        ),
        extracted_text_path=(
            extracted_text_path
            .relative_to(output_root)
            .as_posix()
        ),
        content_sha256=(
            _sha256_bytes(raw_bytes)
        ),
        document_text=document_text,
        page_texts=page_texts,
    )